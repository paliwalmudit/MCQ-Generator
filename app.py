import streamlit as st
import os
import json
import pandas as pd
import requests
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from model import generate_mcqs
import time

#Web page
def get_webpage_content(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    return soup.get_text()

# Streamlit UI setup
st.set_page_config(page_title="MCQ Generator team C4", layout="wide")

st.title("MCQ Generator")

# Sidebar controls
st.sidebar.header("Controls")
number_of_questions = st.sidebar.number_input("Number of questions", min_value=1, value=5)
tone = st.sidebar.selectbox("Tone", options=["simple", "neutral", "professional", "research", "examination"])

# Main area with tabs
tab1, tab2 = st.tabs(["Upload File", "Enter URL"])

with tab1:
    uploaded_file = st.file_uploader("Choose a text file", type="txt")

with tab2:
    webpage_url = st.text_input("Enter a webpage URL")

if st.sidebar.button("Generate MCQs"):
    if uploaded_file is not None:
        text = uploaded_file.getvalue().decode("utf-8")
    elif webpage_url and uploaded_file:
        st.error("Kindly use single type to resolve the ambiguity, use either Document file or a url.")
    elif webpage_url:
        text = get_webpage_content(webpage_url)
    else:
        st.error("Please upload a file or enter a webpage URL")
        st.stop()
    
    with st.spinner("Generating MCQs..."):
        quiz_data, answer_key = generate_mcqs(text, number_of_questions, tone)
        
        if quiz_data:
            st.session_state.quiz_data = quiz_data  # Store generated MCQs in session state
            st.session_state.answer_key = answer_key  # Store answer key in session state
            
            st.subheader("Generated MCQs")
            
            # Display the generated MCQs
            for mcq in quiz_data:
                st.write(f"**{mcq['Question']}**")
                for option, text in mcq['Options'].items():
                    st.write(f"{option}) {text}")
                st.write("---")
        else:
            st.error("No MCQs generated.")
    
# Download options
if 'quiz_data' in st.session_state:
    download_format = st.selectbox("Download as", options=["None","PDF", "Word", "Answer Key"])
    
    if download_format == "PDF":
        pdf_file = SimpleDocTemplate('mcqs.pdf', pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()  # Get default styles
        
        for mcq in st.session_state.quiz_data:
            elements.append(Paragraph(mcq['Question'], styles['Heading1']))
            for option, text in mcq['Options'].items():
                elements.append(Paragraph(f"{option}) {text}", styles['Normal']))
            elements.append(Paragraph("", styles['Normal']))  # Add space between questions
        
        pdf_file.build(elements)
        st.download_button(
            label="Download MCQs as PDF",
            data=open('mcqs.pdf', 'rb').read(),
            file_name="mcqs.pdf",
            mime="application/pdf",
        )
    
    elif download_format == "Word":
        document = Document()
        
        for mcq in st.session_state.quiz_data:
            document.add_heading(mcq['Question'], level=1)
            for option, text in mcq['Options'].items():
                document.add_paragraph(f"{option}) {text}")
            document.add_paragraph()
        
        document.save('mcqs.docx')
        st.download_button(
            label="Download MCQs as Word",
            data=open('mcqs.docx', 'rb').read(),
            file_name="mcqs.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    
    elif download_format == "Answer Key":
        answer_key_doc = Document()
        answer_key_doc.add_heading("Answer Key", level=1)
        
        for question, answer in st.session_state.answer_key.items():
            answer_key_doc.add_paragraph(f"{question} - Correct answer: {answer}")
        
        answer_key_doc.save('answer_key.docx')
        st.download_button(
            label="Download Answer Key",
            data=open('answer_key.docx', 'rb').read(),
            file_name="answer_key.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.write("Select any option to download")
    

st.sidebar.markdown("---")
st.sidebar.header("About")
st.sidebar.info("This app generates multiple-choice questions (MCQs) based on the provided text or webpage content using Google's Gemini AI. Upload a text file or enter a webpage URL, specify the number of questions and tone, then click 'Generate MCQs' to create your quiz.")
